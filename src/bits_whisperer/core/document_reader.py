"""Document reader for extracting text from external files.

Supports reading text from common document formats so they can be
attached as supplementary context to AI actions.  Each reader returns
plain-text content suitable for injection into an LLM prompt.

Supported formats
-----------------
- **Plain text** (.txt, .md, .csv, .log, .json, .xml, .yaml, .yml, .ini, .cfg, .conf, .rst)
- **Word documents** (.docx) — paragraphs + table cells via ``python-docx``
- **Spreadsheets** (.xlsx, .xls) — sheet-by-sheet CSV via ``openpyxl``
- **PDF** (.pdf) — page-by-page text via ``PyPDF2`` / ``pypdf``
- **Rich text** (.rtf) — via ``striprtf``
"""

from __future__ import annotations

import csv
import io
import logging
from pathlib import Path
from typing import Final

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Supported extensions
# ---------------------------------------------------------------------------

_TEXT_EXTENSIONS: Final[frozenset[str]] = frozenset(
    {
        ".txt",
        ".md",
        ".csv",
        ".log",
        ".json",
        ".xml",
        ".yaml",
        ".yml",
        ".ini",
        ".cfg",
        ".conf",
        ".rst",
        ".tsv",
        ".html",
        ".htm",
    }
)

_WORD_EXTENSIONS: Final[frozenset[str]] = frozenset({".docx"})
_EXCEL_EXTENSIONS: Final[frozenset[str]] = frozenset({".xlsx", ".xls"})
_PDF_EXTENSIONS: Final[frozenset[str]] = frozenset({".pdf"})
_RTF_EXTENSIONS: Final[frozenset[str]] = frozenset({".rtf"})

SUPPORTED_EXTENSIONS: Final[frozenset[str]] = (
    _TEXT_EXTENSIONS | _WORD_EXTENSIONS | _EXCEL_EXTENSIONS | _PDF_EXTENSIONS | _RTF_EXTENSIONS
)

# File dialog wildcard for supported attachment types
ATTACHMENT_WILDCARD: Final[str] = (
    "All supported files|"
    "*.txt;*.md;*.csv;*.log;*.json;*.xml;*.yaml;*.yml;*.ini;*.cfg;*.conf;*.rst;"
    "*.tsv;*.html;*.htm;*.docx;*.xlsx;*.xls;*.pdf;*.rtf|"
    "Text files (*.txt, *.md, *.csv, *.log)|*.txt;*.md;*.csv;*.log;*.tsv|"
    "Word documents (*.docx)|*.docx|"
    "Spreadsheets (*.xlsx, *.xls)|*.xlsx;*.xls|"
    "PDF files (*.pdf)|*.pdf|"
    "Rich text (*.rtf)|*.rtf|"
    "Data files (*.json, *.xml, *.yaml)|*.json;*.xml;*.yaml;*.yml|"
    "All files (*.*)|*.*"
)

# Maximum file size to read (10 MB)
_MAX_FILE_SIZE: Final[int] = 10 * 1024 * 1024


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def is_supported(path: str | Path) -> bool:
    """Check whether the file extension is supported for text extraction.

    Args:
        path: File path to check.

    Returns:
        True if the file type is supported.
    """
    return Path(path).suffix.lower() in SUPPORTED_EXTENSIONS


def read_document(path: str | Path) -> str:
    """Read a document and return its text content.

    Dispatches to format-specific readers based on file extension.
    Returns an error message string (prefixed with ``[Error]``) if
    reading fails, rather than raising.

    Args:
        path: Absolute or relative path to the document file.

    Returns:
        Extracted plain-text content from the file.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file is too large (> 10 MB).
    """
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(f"File not found: {p}")

    size = p.stat().st_size
    if size > _MAX_FILE_SIZE:
        raise ValueError(
            f"File is too large ({size / 1024 / 1024:.1f} MB). "
            f"Maximum supported size is {_MAX_FILE_SIZE / 1024 / 1024:.0f} MB."
        )

    ext = p.suffix.lower()

    if ext in _TEXT_EXTENSIONS:
        return _read_text(p)
    elif ext in _WORD_EXTENSIONS:
        return _read_docx(p)
    elif ext in _EXCEL_EXTENSIONS:
        return _read_excel(p)
    elif ext in _PDF_EXTENSIONS:
        return _read_pdf(p)
    elif ext in _RTF_EXTENSIONS:
        return _read_rtf(p)
    else:
        # Attempt plain-text read for unknown extensions
        logger.info("Unknown extension '%s', attempting plain-text read", ext)
        return _read_text(p)


def read_document_safe(path: str | Path) -> str:
    """Read a document, returning an error string on failure instead of raising.

    Args:
        path: Path to the document.

    Returns:
        Extracted text, or a bracketed error message on failure.
    """
    try:
        return read_document(path)
    except Exception as exc:
        return f"[Error reading {Path(path).name}: {exc}]"


# ---------------------------------------------------------------------------
# Format readers
# ---------------------------------------------------------------------------


def _read_text(path: Path) -> str:
    """Read plain text file with encoding detection fallback."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _read_docx(path: Path) -> str:
    """Read text from a .docx Word document.

    Extracts paragraph text and table cell text.
    Requires ``python-docx`` to be installed.
    """
    try:
        from docx import Document
    except ImportError:
        return (
            f"[Cannot read {path.name}: python-docx is not installed. "
            f"Install it with: pip install python-docx]"
        )

    doc = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts)


def _read_excel(path: Path) -> str:
    """Read spreadsheet data as CSV-like text.

    Returns each sheet with a header, and rows as pipe-separated values.
    Requires ``openpyxl`` to be installed.
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        return (
            f"[Cannot read {path.name}: openpyxl is not installed. "
            f"Install it with: pip install openpyxl]"
        )

    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        buf = io.StringIO()
        writer = csv.writer(buf)
        for row_count, row in enumerate(ws.iter_rows(values_only=True)):
            writer.writerow([str(v) if v is not None else "" for v in row])
            if row_count + 1 >= 5000:
                buf.write("... (truncated at 5000 rows)\n")
                break
        sheet_text = buf.getvalue().strip()
        if sheet_text:
            parts.append(f"=== Sheet: {sheet_name} ===\n{sheet_text}")

    wb.close()
    return "\n\n".join(parts)


def _read_pdf(path: Path) -> str:
    """Read text from a PDF file.

    Tries ``pypdf`` first (modern), falls back to ``PyPDF2``.
    """
    # Try pypdf (modern)
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages: list[str] = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"--- Page {i + 1} ---\n{text.strip()}")
        return "\n\n".join(pages) if pages else "[No extractable text in PDF]"
    except ImportError:
        pass

    # Try PyPDF2 (legacy)
    try:
        from PyPDF2 import PdfReader as LegacyReader

        reader = LegacyReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"--- Page {i + 1} ---\n{text.strip()}")
        return "\n\n".join(pages) if pages else "[No extractable text in PDF]"
    except ImportError:
        return (
            f"[Cannot read {path.name}: pypdf is not installed. "
            f"Install it with: pip install pypdf]"
        )


def _read_rtf(path: Path) -> str:
    """Read text from an RTF file.

    Requires ``striprtf`` to be installed.
    """
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError:
        return (
            f"[Cannot read {path.name}: striprtf is not installed. "
            f"Install it with: pip install striprtf]"
        )

    raw = path.read_bytes().decode("utf-8", errors="replace")
    return str(rtf_to_text(raw))
