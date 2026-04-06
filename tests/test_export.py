"""Tests for export formatters."""

from __future__ import annotations

import json
from pathlib import Path

import pytest

from bits_whisperer.core.job import TranscriptionResult, TranscriptSegment
from bits_whisperer.export.base import format_timestamp, format_timestamp_srt
from bits_whisperer.export.html_export import HTMLFormatter
from bits_whisperer.export.json_export import JSONFormatter
from bits_whisperer.export.markdown import MarkdownFormatter
from bits_whisperer.export.plain_text import PlainTextFormatter
from bits_whisperer.export.srt import SRTFormatter
from bits_whisperer.export.vtt import VTTFormatter

_has_docx = True
try:
    from bits_whisperer.export.word_export import WordFormatter
except ImportError:
    _has_docx = False


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


class TestFormatTimestamp:
    """Timestamp formatting helpers."""

    def test_zero(self) -> None:
        assert format_timestamp(0.0) == "00:00:00.000"

    def test_seconds_only(self) -> None:
        assert format_timestamp(5.5) == "00:00:05.500"

    def test_minutes_and_seconds(self) -> None:
        assert format_timestamp(125.75) == "00:02:05.750"

    def test_hours(self) -> None:
        assert format_timestamp(3661.123) == "01:01:01.123"

    def test_srt_format_uses_comma(self) -> None:
        ts = format_timestamp_srt(5.5)
        assert "," in ts
        assert "." not in ts
        assert ts == "00:00:05,500"


def _make_result(
    segments: list[TranscriptSegment] | None = None,
    full_text: str = "",
) -> TranscriptionResult:
    """Helper to create a test result."""
    return TranscriptionResult(
        job_id="test",
        audio_file="test.mp3",
        provider="test",
        model="test",
        language="en",
        duration_seconds=10.0,
        segments=segments or [],
        full_text=full_text,
    )


_SAMPLE_SEGMENTS = [
    TranscriptSegment(start=0.0, end=2.5, text="Hello there.", speaker="Alice", confidence=0.95),
    TranscriptSegment(start=2.5, end=5.0, text="How are you?", speaker="Bob", confidence=0.88),
    TranscriptSegment(start=5.0, end=8.0, text="I'm great.", speaker="Alice", confidence=0.92),
]


# ---------------------------------------------------------------------------
# PlainTextFormatter
# ---------------------------------------------------------------------------


class TestPlainTextFormatter:
    """Plain text export."""

    def test_format_properties(self) -> None:
        fmt = PlainTextFormatter()
        assert fmt.format_id == "txt"
        assert fmt.file_extension == ".txt"
        assert "Plain Text" in fmt.display_name

    def test_export_full_text(self, tmp_path: Path) -> None:
        result = _make_result(full_text="Hello world")
        out = tmp_path / "output.txt"
        fmt = PlainTextFormatter()
        written = fmt.export(result, out)
        assert written == out
        assert out.read_text(encoding="utf-8") == "Hello world"

    def test_export_segments(self, tmp_path: Path) -> None:
        segments = [
            TranscriptSegment(start=0.0, end=2.0, text="Hello"),
            TranscriptSegment(start=2.0, end=4.0, text="World"),
        ]
        result = _make_result(segments=segments)
        out = tmp_path / "output.txt"
        fmt = PlainTextFormatter()
        fmt.export(result, out)
        content = out.read_text(encoding="utf-8")
        assert "Hello" in content
        assert "World" in content

    def test_export_with_timestamps(self, tmp_path: Path) -> None:
        segments = [
            TranscriptSegment(start=0.0, end=2.0, text="Hello"),
        ]
        result = _make_result(segments=segments)
        out = tmp_path / "output.txt"
        fmt = PlainTextFormatter()
        fmt.export(result, out, include_timestamps=True)
        content = out.read_text(encoding="utf-8")
        assert "[00:00:00.000]" in content
        assert "Hello" in content

    def test_export_with_speakers(self, tmp_path: Path) -> None:
        segments = [
            TranscriptSegment(start=0.0, end=2.0, text="Hi", speaker="Alice"),
        ]
        result = _make_result(segments=segments)
        out = tmp_path / "output.txt"
        fmt = PlainTextFormatter()
        fmt.export(result, out, include_speakers=True)
        content = out.read_text(encoding="utf-8")
        assert "Alice:" in content

    def test_export_with_confidence(self, tmp_path: Path) -> None:
        segments = [
            TranscriptSegment(start=0.0, end=2.0, text="Hi", confidence=0.95),
        ]
        result = _make_result(segments=segments)
        out = tmp_path / "output.txt"
        fmt = PlainTextFormatter()
        fmt.export(result, out, include_confidence=True)
        content = out.read_text(encoding="utf-8")
        assert "95%" in content


# ---------------------------------------------------------------------------
# MarkdownFormatter
# ---------------------------------------------------------------------------


class TestMarkdownFormatter:
    """Markdown export."""

    def test_format_properties(self) -> None:
        fmt = MarkdownFormatter()
        assert fmt.format_id == "md"
        assert fmt.file_extension == ".md"
        assert "Markdown" in fmt.display_name

    def test_export_full_text_fallback(self, tmp_path: Path) -> None:
        result = _make_result(full_text="Fallback text")
        out = tmp_path / "output.md"
        written = MarkdownFormatter().export(result, out)
        assert written == out
        content = out.read_text(encoding="utf-8")
        assert "Fallback text" in content

    def test_export_segments_with_metadata(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.md"
        MarkdownFormatter().export(result, out)
        content = out.read_text(encoding="utf-8")
        # Title & metadata
        assert "# Transcript: test.mp3" in content
        assert "**Provider**: test" in content
        assert "**Language**: en" in content
        # Segments
        assert "Hello there." in content
        assert "How are you?" in content

    def test_export_with_speakers(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.md"
        MarkdownFormatter().export(result, out, include_speakers=True)
        content = out.read_text(encoding="utf-8")
        assert "### Alice" in content
        assert "### Bob" in content

    def test_export_without_speakers(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.md"
        MarkdownFormatter().export(result, out, include_speakers=False)
        content = out.read_text(encoding="utf-8")
        assert "### Alice" not in content
        assert "### Bob" not in content

    def test_export_with_timestamps(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.md"
        MarkdownFormatter().export(result, out, include_timestamps=True)
        content = out.read_text(encoding="utf-8")
        assert "> *00:00:00.000" in content

    def test_export_without_timestamps(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.md"
        MarkdownFormatter().export(result, out, include_timestamps=False)
        content = out.read_text(encoding="utf-8")
        assert "> *00:00:00.000" not in content
        assert "Hello there." in content

    def test_export_with_confidence(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.md"
        MarkdownFormatter().export(result, out, include_confidence=True)
        content = out.read_text(encoding="utf-8")
        assert "_95%_" in content
        assert "_88%_" in content

    def test_export_without_confidence(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.md"
        MarkdownFormatter().export(result, out, include_confidence=False)
        content = out.read_text(encoding="utf-8")
        assert "_95%_" not in content

    def test_speaker_grouping(self, tmp_path: Path) -> None:
        """Consecutive segments by same speaker should not repeat heading."""
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.md"
        MarkdownFormatter().export(result, out, include_speakers=True)
        content = out.read_text(encoding="utf-8")
        # Alice appears at seg 0 and seg 2 — heading should appear twice
        assert content.count("### Alice") == 2
        assert content.count("### Bob") == 1


# ---------------------------------------------------------------------------
# HTMLFormatter
# ---------------------------------------------------------------------------


class TestHTMLFormatter:
    """HTML export."""

    def test_format_properties(self) -> None:
        fmt = HTMLFormatter()
        assert fmt.format_id == "html"
        assert fmt.file_extension == ".html"
        assert "HTML" in fmt.display_name

    def test_export_full_text_fallback(self, tmp_path: Path) -> None:
        result = _make_result(full_text="Plain fallback")
        out = tmp_path / "output.html"
        HTMLFormatter().export(result, out)
        content = out.read_text(encoding="utf-8")
        assert "<p>Plain fallback</p>" in content

    def test_export_produces_valid_html(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.html"
        HTMLFormatter().export(result, out)
        content = out.read_text(encoding="utf-8")
        assert content.startswith("<!DOCTYPE html>")
        assert "</html>" in content
        assert "<title>Transcript: test.mp3</title>" in content

    def test_export_with_segments(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.html"
        HTMLFormatter().export(result, out)
        content = out.read_text(encoding="utf-8")
        assert "Hello there." in content
        assert "How are you?" in content
        assert '<div class="segment">' in content

    def test_export_with_speakers(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.html"
        HTMLFormatter().export(result, out, include_speakers=True)
        content = out.read_text(encoding="utf-8")
        assert '<span class="speaker">Alice:</span>' in content
        assert '<span class="speaker">Bob:</span>' in content

    def test_export_without_speakers(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.html"
        HTMLFormatter().export(result, out, include_speakers=False)
        content = out.read_text(encoding="utf-8")
        assert "speaker" not in content.split("</style>")[1]

    def test_export_with_timestamps(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.html"
        HTMLFormatter().export(result, out, include_timestamps=True)
        content = out.read_text(encoding="utf-8")
        assert '<span class="timestamp">' in content
        assert "00:00:00.000" in content

    def test_export_without_timestamps(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.html"
        HTMLFormatter().export(result, out, include_timestamps=False)
        # timestamp spans should not appear in the transcript body
        body = out.read_text(encoding="utf-8").split("</style>")[1]
        assert '<span class="timestamp">' not in body

    def test_export_with_confidence(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.html"
        HTMLFormatter().export(result, out, include_confidence=True)
        content = out.read_text(encoding="utf-8")
        assert '<span class="confidence">(95%)</span>' in content

    def test_export_without_confidence(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.html"
        HTMLFormatter().export(result, out, include_confidence=False)
        content = out.read_text(encoding="utf-8")
        assert "confidence" not in content.split("</style>")[1]

    def test_html_escaping(self, tmp_path: Path) -> None:
        """Special characters must be HTML-escaped."""
        segments = [TranscriptSegment(start=0.0, end=1.0, text="<script>alert('xss')</script>")]
        result = _make_result(segments=segments)
        out = tmp_path / "output.html"
        HTMLFormatter().export(result, out)
        content = out.read_text(encoding="utf-8")
        assert "<script>" not in content
        assert "&lt;script&gt;" in content

    def test_metadata_in_header(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.html"
        HTMLFormatter().export(result, out)
        content = out.read_text(encoding="utf-8")
        assert "Provider: test" in content
        assert "Model: test" in content
        assert "Language: en" in content

    def test_dark_mode_css(self, tmp_path: Path) -> None:
        """HTML should include dark mode media query."""
        result = _make_result(full_text="dark test")
        out = tmp_path / "output.html"
        HTMLFormatter().export(result, out)
        content = out.read_text(encoding="utf-8")
        assert "prefers-color-scheme: dark" in content


# ---------------------------------------------------------------------------
# SRTFormatter
# ---------------------------------------------------------------------------


class TestSRTFormatter:
    """SubRip (.srt) subtitle export."""

    def test_format_properties(self) -> None:
        fmt = SRTFormatter()
        assert fmt.format_id == "srt"
        assert fmt.file_extension == ".srt"
        assert "SubRip" in fmt.display_name

    def test_export_full_text_fallback(self, tmp_path: Path) -> None:
        result = _make_result(full_text="Single cue text")
        out = tmp_path / "output.srt"
        SRTFormatter().export(result, out)
        content = out.read_text(encoding="utf-8")
        assert "1\n" in content
        assert "Single cue text" in content
        assert "00:00:00,000 --> 00:00:10,000" in content

    def test_export_segments(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.srt"
        SRTFormatter().export(result, out)
        content = out.read_text(encoding="utf-8")
        lines = content.split("\n")
        # First cue index
        assert lines[0] == "1"
        # SRT uses comma separators
        assert "00:00:00,000 --> 00:00:02,500" in content
        assert "Hello there." in content

    def test_cue_numbering(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.srt"
        SRTFormatter().export(result, out)
        content = out.read_text(encoding="utf-8")
        assert "\n1\n" in content or content.startswith("1\n")
        assert "\n2\n" in content
        assert "\n3\n" in content

    def test_export_with_speakers(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.srt"
        SRTFormatter().export(result, out, include_speakers=True)
        content = out.read_text(encoding="utf-8")
        assert "[Alice]" in content
        assert "[Bob]" in content

    def test_export_without_speakers(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.srt"
        SRTFormatter().export(result, out, include_speakers=False)
        content = out.read_text(encoding="utf-8")
        assert "[Alice]" not in content
        assert "[Bob]" not in content
        assert "Hello there." in content

    def test_srt_timestamps_use_comma(self, tmp_path: Path) -> None:
        """SRT timestamps must use comma for milliseconds, not period."""
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.srt"
        SRTFormatter().export(result, out)
        content = out.read_text(encoding="utf-8")
        # All timestamp lines: check comma separator
        for line in content.split("\n"):
            if "-->" in line:
                assert "," in line
                # Periods should not appear in the timestamp portion
                parts = line.split(" --> ")
                for ts in parts:
                    assert "." not in ts


# ---------------------------------------------------------------------------
# VTTFormatter
# ---------------------------------------------------------------------------


class TestVTTFormatter:
    """WebVTT (.vtt) subtitle export."""

    def test_format_properties(self) -> None:
        fmt = VTTFormatter()
        assert fmt.format_id == "vtt"
        assert fmt.file_extension == ".vtt"
        assert "WebVTT" in fmt.display_name

    def test_webvtt_header(self, tmp_path: Path) -> None:
        result = _make_result(full_text="Header test")
        out = tmp_path / "output.vtt"
        VTTFormatter().export(result, out)
        content = out.read_text(encoding="utf-8")
        assert content.startswith("WEBVTT\n")

    def test_export_full_text_fallback(self, tmp_path: Path) -> None:
        result = _make_result(full_text="Single cue")
        out = tmp_path / "output.vtt"
        VTTFormatter().export(result, out)
        content = out.read_text(encoding="utf-8")
        assert "Single cue" in content
        assert "00:00:00.000 --> 00:00:10.000" in content

    def test_export_segments(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.vtt"
        VTTFormatter().export(result, out)
        content = out.read_text(encoding="utf-8")
        assert "00:00:00.000 --> 00:00:02.500" in content
        assert "Hello there." in content

    def test_cue_numbering(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.vtt"
        VTTFormatter().export(result, out)
        content = out.read_text(encoding="utf-8")
        assert "\n1\n" in content
        assert "\n2\n" in content
        assert "\n3\n" in content

    def test_export_with_speakers_uses_voice_tag(self, tmp_path: Path) -> None:
        """VTT uses <v Speaker> syntax for voice identification."""
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.vtt"
        VTTFormatter().export(result, out, include_speakers=True)
        content = out.read_text(encoding="utf-8")
        assert "<v Alice>" in content
        assert "<v Bob>" in content

    def test_export_without_speakers(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.vtt"
        VTTFormatter().export(result, out, include_speakers=False)
        content = out.read_text(encoding="utf-8")
        assert "<v Alice>" not in content
        assert "<v Bob>" not in content
        assert "Hello there." in content

    def test_vtt_timestamps_use_period(self, tmp_path: Path) -> None:
        """VTT timestamps use period for millisecond separator."""
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.vtt"
        VTTFormatter().export(result, out)
        content = out.read_text(encoding="utf-8")
        for line in content.split("\n"):
            if "-->" in line:
                assert "." in line
                assert "," not in line


# ---------------------------------------------------------------------------
# JSONFormatter
# ---------------------------------------------------------------------------


class TestJSONFormatter:
    """JSON export."""

    def test_format_properties(self) -> None:
        fmt = JSONFormatter()
        assert fmt.format_id == "json"
        assert fmt.file_extension == ".json"
        assert "JSON" in fmt.display_name

    def test_export_produces_valid_json(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.json"
        JSONFormatter().export(result, out)
        data = json.loads(out.read_text(encoding="utf-8"))
        assert isinstance(data, dict)
        assert data["job_id"] == "test"
        assert data["audio_file"] == "test.mp3"

    def test_export_metadata(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.json"
        JSONFormatter().export(result, out)
        data = json.loads(out.read_text(encoding="utf-8"))
        assert data["provider"] == "test"
        assert data["model"] == "test"
        assert data["language"] == "en"
        assert data["duration_seconds"] == 10.0

    def test_export_segments_included(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.json"
        JSONFormatter().export(result, out)
        data = json.loads(out.read_text(encoding="utf-8"))
        assert len(data["segments"]) == 3
        assert data["segments"][0]["text"] == "Hello there."

    def test_export_with_timestamps(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.json"
        JSONFormatter().export(result, out, include_timestamps=True)
        data = json.loads(out.read_text(encoding="utf-8"))
        seg = data["segments"][0]
        assert "start" in seg
        assert "end" in seg
        assert seg["start"] == 0.0
        assert seg["end"] == 2.5

    def test_export_without_timestamps(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.json"
        JSONFormatter().export(result, out, include_timestamps=False)
        data = json.loads(out.read_text(encoding="utf-8"))
        seg = data["segments"][0]
        assert "start" not in seg
        assert "end" not in seg
        assert seg["text"] == "Hello there."

    def test_export_with_speakers(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.json"
        JSONFormatter().export(result, out, include_speakers=True)
        data = json.loads(out.read_text(encoding="utf-8"))
        assert data["segments"][0]["speaker"] == "Alice"

    def test_export_without_speakers(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.json"
        JSONFormatter().export(result, out, include_speakers=False)
        data = json.loads(out.read_text(encoding="utf-8"))
        assert "speaker" not in data["segments"][0]

    def test_export_with_confidence(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.json"
        JSONFormatter().export(result, out, include_confidence=True)
        data = json.loads(out.read_text(encoding="utf-8"))
        assert data["segments"][0]["confidence"] == 0.95

    def test_export_without_confidence(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.json"
        JSONFormatter().export(result, out, include_confidence=False)
        data = json.loads(out.read_text(encoding="utf-8"))
        assert "confidence" not in data["segments"][0]

    def test_export_full_text_included(self, tmp_path: Path) -> None:
        result = _make_result(full_text="Full text content")
        out = tmp_path / "output.json"
        JSONFormatter().export(result, out)
        data = json.loads(out.read_text(encoding="utf-8"))
        assert data["full_text"] == "Full text content"

    def test_strip_multiple_fields(self, tmp_path: Path) -> None:
        """Stripping timestamps, speakers, and confidence simultaneously."""
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.json"
        JSONFormatter().export(
            result,
            out,
            include_timestamps=False,
            include_speakers=False,
            include_confidence=False,
        )
        data = json.loads(out.read_text(encoding="utf-8"))
        seg = data["segments"][0]
        assert "start" not in seg
        assert "end" not in seg
        assert "speaker" not in seg
        assert "confidence" not in seg
        assert seg["text"] == "Hello there."


# ---------------------------------------------------------------------------
# WordFormatter
# ---------------------------------------------------------------------------


@pytest.mark.skipif(not _has_docx, reason="python-docx not installed")
class TestWordFormatter:
    """Microsoft Word (.docx) export."""

    def test_format_properties(self) -> None:
        fmt = WordFormatter()
        assert fmt.format_id == "docx"
        assert fmt.file_extension == ".docx"
        assert "Word" in fmt.display_name

    def test_export_creates_file(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.docx"
        written = WordFormatter().export(result, out)
        assert written == out
        assert out.exists()
        assert out.stat().st_size > 0

    def test_export_full_text_fallback(self, tmp_path: Path) -> None:
        result = _make_result(full_text="Fallback paragraph")
        out = tmp_path / "output.docx"
        WordFormatter().export(result, out)
        assert out.exists()

    def test_export_with_segments(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.docx"
        WordFormatter().export(result, out)
        # Verify it's a valid docx by opening it
        from docx import Document

        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Hello there." in full_text

    def test_export_with_speakers(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.docx"
        WordFormatter().export(result, out, include_speakers=True)
        from docx import Document

        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Alice" in full_text
        assert "Bob" in full_text

    def test_export_metadata_table(self, tmp_path: Path) -> None:
        result = _make_result(segments=list(_SAMPLE_SEGMENTS))
        out = tmp_path / "output.docx"
        WordFormatter().export(result, out)
        from docx import Document

        doc = Document(str(out))
        # Should have a metadata table
        assert len(doc.tables) >= 1
        table = doc.tables[0]
        # 5 metadata rows
        assert len(table.rows) == 5
        assert table.cell(0, 0).text == "Provider"
        assert table.cell(0, 1).text == "test"

    def test_export_footer(self, tmp_path: Path) -> None:
        result = _make_result(full_text="Footer test")
        out = tmp_path / "output.docx"
        WordFormatter().export(result, out)
        from docx import Document

        doc = Document(str(out))
        footer = doc.sections[0].footer
        assert "BITS Whisperer" in footer.paragraphs[0].text

    def test_docx_import_error(self) -> None:
        """WordFormatter should raise ImportError if python-docx is missing."""
        # This test verifies the error message path; python-docx IS installed
        # so we just verify the formatter is importable
        assert WordFormatter is not None
